import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import random
import csv
from collections import defaultdict
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import os
import platform
import threading

class ExamInfoForm(tk.Toplevel):
    def __init__(self, master, save_callback):
        super().__init__(master)
        self.title("Exam Information")
        self.geometry("500x400")

        self.save_callback = save_callback
        labels = [
            "Institute Name", "Exam Title", "Program Name",
            "Course/Paper Name", "Exam Time", "Exam Date", "Total Marks"
        ]
        self.entries = {}
        for i, label in enumerate(labels):
            tk.Label(self, text=label + ":").grid(row=i, column=0, sticky='e', padx=10, pady=5)
            entry = tk.Entry(self, width=40)
            entry.grid(row=i, column=1, padx=10, pady=5)
            self.entries[label] = entry

        tk.Button(self, text="Save", command=self.save).grid(row=len(labels), column=0, columnspan=2, pady=20)

    def save(self):
        data = {label: self.entries[label].get().strip() for label in self.entries}
        if all(data.values()):
            self.save_callback(data)
            self.destroy()
        else:
            messagebox.showerror("Missing Info", "Please fill in all fields.")


class QuestionGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Bloom's Taxonomy Question Generator")
        self.exam_info = {}
        self.question_bank = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
        self.levels = ["Remembering", "Understanding", "Applying", "Analyzing", "Evaluating", "Creating"]
        self.inputs = {}
        self.subjects = set()
        self.selected_subject = tk.StringVar(value="All")
        self.generated_q1 = []
        self.generated_q2 = []
        self.build_gui()

    def build_gui(self):
        control_frame = tk.Frame(self.root)
        control_frame.pack(pady=10, fill='x')

        tk.Button(control_frame, text="Enter Exam Info", command=self.open_exam_info_form).pack(side='left', padx=5)
        tk.Button(control_frame, text="Load CSV", command=self.load_csv_dialog).pack(side='left', padx=5)

        tk.Label(control_frame, text="Select Subject:").pack(side='left', padx=5)
        self.subject_combo = ttk.Combobox(control_frame, values=["All"], state="readonly", textvariable=self.selected_subject)
        self.subject_combo.pack(side='left', padx=5)
        self.subject_combo.bind("<<ComboboxSelected>>", lambda e: self.update_spinboxes())

        level_frame = tk.Frame(self.root)
        level_frame.pack(pady=10)
        for level in self.levels:
            frame = tk.Frame(level_frame)
            frame.pack(pady=2)
            tk.Label(frame, text=level, width=15).pack(side='left')
            var = tk.IntVar(value=1)
            self.inputs[level] = var
            tk.Spinbox(frame, from_=0, to=10, width=5, textvariable=var).pack(side='left')

        tk.Button(self.root, text="Generate Questions", command=self.generate_questions).pack(pady=10)

        self.output = tk.Text(self.root, wrap=tk.WORD, height=20, width=100)
        self.output.pack(padx=10, pady=10)

        export_frame = tk.Frame(self.root)
        export_frame.pack(pady=10)
        tk.Button(export_frame, text="Export to PDF", command=self.export_pdf).pack(side='left', padx=5)
        tk.Button(export_frame, text="Export to Word", command=self.export_word).pack(side='left', padx=5)

    def open_exam_info_form(self):
        ExamInfoForm(self.root, self.save_exam_info)

    def save_exam_info(self, info):
        self.exam_info = info
        messagebox.showinfo("Saved", "Exam information saved successfully.")

    def load_csv_dialog(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
        if not file_path:
            return
        threading.Thread(target=self.load_questions_from_csv, args=(file_path,), daemon=True).start()

    def load_questions_from_csv(self, file_path):
        import chardet
        self.question_bank.clear()
        self.subjects.clear()
        try:
            with open(file_path, 'rb') as f:
                encoding = chardet.detect(f.read())['encoding']
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    subject = row['Subject'].strip()
                    level = row['Level'].strip()
                    marks = row.get('Marks', '1').strip()
                    question = row['Question'].strip()
                    co = row.get('CO', '').strip()
                    cl = row.get('CL', '').strip()
                    kc = row.get('KC', '').strip()
                    if level in self.levels:
                        self.question_bank[subject][level][marks].append((question, co, cl, kc))
                        self.subjects.add(subject)
            subjects = sorted(list(self.subjects))
            self.subject_combo['values'] = ["All"] + subjects
            self.subject_combo.set("All")
            messagebox.showinfo("Loaded", f"Loaded questions successfully.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load CSV: {e}")

    def update_spinboxes(self):
        pass

    def generate_questions(self):
        self.output.delete(1.0, tk.END)
        selected = self.selected_subject.get()
        if selected == "All":
            data = defaultdict(lambda: defaultdict(list))
            for subject in self.subjects:
                for level in self.levels:
                    for marks, items in self.question_bank[subject][level].items():
                        data[level][marks].extend(items)
        else:
            data = self.question_bank[selected]

        self.generated_q1.clear()
        self.generated_q2.clear()

        for level in ["Remembering", "Understanding"]:
            available = data.get(level, {}).get('1', [])
            count = self.inputs[level].get()
            if count > len(available):
                count = len(available)
            self.generated_q1.extend(random.sample(available, count))

        for level in ["Applying", "Analyzing", "Evaluating", "Creating"]:
            available = data.get(level, {}).get('5', [])
            count = self.inputs[level].get()
            if count > len(available):
                count = len(available)
            self.generated_q2.extend(random.sample(available, count))

        self.output.insert(tk.END, "Q-1 Answer all questions.\n")
        for i, (q, co, cl, kc) in enumerate(self.generated_q1, 1):
            self.output.insert(tk.END, f"{i}. {q} (CO: {co}, CL: {cl}, KC: {kc})\n")
        self.output.insert(tk.END, "\nQ-2 Answer any four out of five questions.\n")
        for i, (q, co, cl, kc) in enumerate(self.generated_q2, 1):
            self.output.insert(tk.END, f"{i}. {q} (CO: {co}, CL: {cl}, KC: {kc})\n")

    def export_pdf(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if not file_path:
            return
        try:
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            elements = []
            for key in ["Institute Name", "Exam Title"]:
                elements.append(Paragraph(self.exam_info.get(key, ''), styles['Title']))
            for key in ["Program Name", "Course/Paper Name"]:
                elements.append(Paragraph(self.exam_info.get(key, ''), styles['Normal']))
            line = f"Time: {self.exam_info.get('Exam Time', '')}    Date: {self.exam_info.get('Exam Date', '')}    Total Marks: {self.exam_info.get('Total Marks', '')}"
            elements.append(Paragraph(line, styles['Normal']))
            elements.append(Spacer(1, 12))

            elements.append(Paragraph("Q-1 Answer all questions. (10×1=10 Marks)", styles['Heading2']))
            data_q1 = [["CO", "CL", "KC", "Question"]] + [[co, cl, kc, q] for q, co, cl, kc in self.generated_q1]
            table_q1 = Table(data_q1, colWidths=[50, 50, 50, 350])
            table_q1.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold')
            ]))
            elements.append(table_q1)
            elements.append(Spacer(1, 12))

            elements.append(Paragraph("Q-2 Answer any four out of five questions. (4×5=20 Marks)", styles['Heading2']))
            data_q2 = [["CO", "CL", "KC", "Question"]] + [[co, cl, kc, q] for q, co, cl, kc in self.generated_q2]
            table_q2 = Table(data_q2, colWidths=[50, 50, 50, 350])
            table_q2.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold')
            ]))
            elements.append(table_q2)
            doc.build(elements)
            messagebox.showinfo("Success", "PDF exported successfully!")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def export_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if not file_path:
            return
        try:
            doc = Document()
            for key in ["Institute Name", "Exam Title"]:
                doc.add_paragraph(self.exam_info.get(key, ''), style='Title')
            for key in ["Program Name", "Course/Paper Name"]:
                doc.add_paragraph(self.exam_info.get(key, ''))
            doc.add_paragraph(f"Time: {self.exam_info.get('Exam Time', '')}    Date: {self.exam_info.get('Exam Date', '')}    Total Marks: {self.exam_info.get('Total Marks', '')}")

            doc.add_paragraph("Q-1 Answer all questions. (10×1=10 Marks)")
            table_q1 = doc.add_table(rows=1, cols=4)
            hdr_cells = table_q1.rows[0].cells
            hdr_cells[0].text = 'CO'
            hdr_cells[1].text = 'CL'
            hdr_cells[2].text = 'KC'
            hdr_cells[3].text = 'Question'
            for q, co, cl, kc in self.generated_q1:
                row = table_q1.add_row().cells
                row[0].text = co
                row[1].text = cl
                row[2].text = kc
                row[3].text = q

            doc.add_paragraph("\nQ-2 Answer any four out of five questions. (4×5=20 Marks)")
            table_q2 = doc.add_table(rows=1, cols=4)
            hdr_cells = table_q2.rows[0].cells
            hdr_cells[0].text = 'CO'
            hdr_cells[1].text = 'CL'
            hdr_cells[2].text = 'KC'
            hdr_cells[3].text = 'Question'
            for q, co, cl, kc in self.generated_q2:
                row = table_q2.add_row().cells
                row[0].text = co
                row[1].text = cl
                row[2].text = kc
                row[3].text = q

            doc.save(file_path)
            messagebox.showinfo("Success", "Word document exported successfully!")
        except Exception as e:
            messagebox.showerror("Error", str(e))


if __name__ == "__main__":
    root = tk.Tk()
    app = QuestionGeneratorApp(root)
    root.mainloop()
